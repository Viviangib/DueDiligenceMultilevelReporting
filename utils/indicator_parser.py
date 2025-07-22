# This file has been split for clarity.
# See utils/file_extraction.py for file extraction functions.
# See utils/indicator_parsing.py for indicator parsing and LLM logic.

import fitz
from docx import Document as DocxDocument
from typing import List
from openai import AsyncOpenAI
from config import settings
from docx import Document as OutputDocx
import io
import json
from utils.prompts.indicator import INDICATOR_PROMPT
import re
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from services.openAI.chat import OpenAIClient


openai_client = OpenAIClient()


def split_text_into_chunks(text: str, chunk_size=3000, chunk_overlap=200) -> list:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", " "],
    )
    return splitter.split_documents([Document(page_content=text)])


def try_extract_json(content: str) -> List[dict]:
    try:
        return json.loads(content)
    except:
        match = re.search(r"\[\s*{[\s\S]*?}\s*]", content)
        if match:
            try:
                return json.loads(match.group())
            except Exception as e:
                print(f"❌ Regex parse failed: {e}")
    return []


def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    print("🔄 Starting PDF text extraction...")
    print(f"📊 PDF file size: {len(file_bytes)} bytes")

    text = ""
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            print(f"📄 PDF has {len(doc)} pages")
            for page_num, page in enumerate(doc, 1):  # type: ignore
                page_text = page.get_text("text")
                text += page_text
                print(
                    f"✅ Extracted text from page {page_num}: {len(page_text)} characters"
                )

        print(f"✅ PDF extraction completed successfully!")
        print(f"📝 Total extracted text length: {len(text)} characters")
        print(f"🔤 First 200 characters: {text[:200]}...")

    except Exception as e:
        print(f"❌ PDF extraction failed: {e}")

    return text


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    print("🔄 Extracting text from DOCX...")
    text = ""
    try:
        file_stream = io.BytesIO(file_bytes)
        doc = DocxDocument(file_stream)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"❌ DOCX extraction failed: {e}")
    return text


async def parse_indicators_with_llm(text: str) -> List[dict]:
    print("🤖 Splitting into chunks for LLM parsing...")
    chunks = split_text_into_chunks(text)
    all_indicators = []

    for i, chunk_doc in enumerate(chunks):
        chunk = chunk_doc.page_content
        prompt = INDICATOR_PROMPT.format(chunk=chunk)

        print(f"📦 Processing chunk {i+1}/{len(chunks)} - {len(chunk)} characters")
        try:

            content = await openai_client.chat(
                prompt=prompt, temperature=0, max_tokens=4000
            )
            print("\\n Content is ", content, "\n\n")

            if content:
                indicators = try_extract_json(content)
                all_indicators.extend(indicators)
            else:

                print(f"⚠️ Empty content in chunk {i+1}, skipping.")

        except Exception as e:
            print(f"❌ Failed on chunk {i+1}: {e}")

    return all_indicators


def save_to_docx(text: str, output_path: str):
    print(f"💾 Saving output to {output_path}...")
    try:
        doc = OutputDocx()
        doc.add_paragraph(text)
        doc.save(output_path)
        print("✅ Save complete")
    except Exception as e:
        print(f"❌ Failed to save DOCX: {e}")
