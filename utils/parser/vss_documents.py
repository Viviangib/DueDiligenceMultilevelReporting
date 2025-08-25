import os
import logging
from typing import List, Tuple
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)

def extract_text_with_pages(file_path: str) -> Tuple[List[str], str]:
    """
    Extract text from a file with page-level information
    
    Args:
        file_path: Path to the file to extract text from
        
    Returns:
        Tuple of (list of page texts, source filename)
    """
    source_file = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        return extract_pdf_with_pages(file_path), source_file
    elif ext == ".docx":
        return extract_docx_with_pages(file_path), source_file
    else:
        raise ValueError(f"Unsupported file type: {ext}")

def extract_pdf_with_pages(file_path: str) -> List[str]:
    """
    Extract text from PDF with page-level information
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        List of text content for each page
    """
    pages_text = []
    
    try:
        with pdfplumber.open(file_path) as pdf:
            logger.info(f"Processing PDF: {file_path}, Total pages: {len(pdf.pages)}")
            
            for page_num, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text.strip())
                else:
                    # Add empty string for pages with no text
                    pages_text.append("")
                    
        logger.info(f"Successfully extracted text from {len(pages_text)} pages")
        return pages_text
        
    except Exception as e:
        logger.error(f"Failed to extract text from PDF {file_path}: {e}")
        raise

def extract_docx_with_pages(file_path: str) -> List[str]:
    """
    Extract text from DOCX file
    
    Note: DOCX doesn't have native page breaks, so we'll simulate pages
    by grouping paragraphs into chunks of roughly page-sized content
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        List of text content for each "page"
    """
    try:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Group paragraphs into page-sized chunks (~500 words per page)
        pages_text = []
        current_page = []
        current_word_count = 0
        words_per_page = 500
        
        for paragraph in paragraphs:
            word_count = len(paragraph.split())
            current_word_count += word_count
            
            if current_word_count > words_per_page and current_page:
                # Start a new page
                pages_text.append("\n".join(current_page))
                current_page = [paragraph]
                current_word_count = word_count
            else:
                current_page.append(paragraph)
        
        # Add the last page
        if current_page:
            pages_text.append("\n".join(current_page))
        
        logger.info(f"Successfully extracted text from DOCX: {file_path}, created {len(pages_text)} page chunks")
        return pages_text
        
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
        raise

def extract_text_simple(file_path: str) -> str:
    """
    Extract text from a file without page-level information (fallback)
    
    Args:
        file_path: Path to the file
        
    Returns:
        Full text content
    """
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            text = "".join(page.extract_text() or "" for page in pdf.pages)
            return text
    elif ext == ".docx":
        doc = Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
        return text
    else:
        raise ValueError(f"Unsupported file type: {ext}") 