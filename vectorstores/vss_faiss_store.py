import os
# Set tokenizers parallelism before importing SentenceTransformer
# This enables parallelism and prevents warnings when processes are forked
os.environ["TOKENIZERS_PARALLELISM"] = "true"

import logging
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
import pdfplumber
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import tiktoken

logger = logging.getLogger(__name__)


@dataclass
class VSSChunk:
    """Represents a chunk of VSS text with metadata"""

    content: str
    page_number: int
    source_file: str
    chunk_id: str
    chunk_index: int


def chunk_text(text: str, chunk_size: int = 200, chunk_overlap: int = 20) -> List[Document]:
    """
    Chunk text into smaller pieces using token-based splitting

    Args:
        text: Text to chunk
        chunk_size: Target size in tokens
        chunk_overlap: Overlap between chunks in tokens

    Returns:
        List of Document objects
    """

    def count_tokens(text: str) -> int:
        try:
            encoding = tiktoken.get_encoding("cl100k_base")
            return len(encoding.encode(text))
        except Exception:
            return len(text.split())  # Fallback to word count

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ".", " ", ""],
        length_function=count_tokens,
    )

    return text_splitter.split_documents([Document(page_content=text)])


class InMemoryVSSVectorStore:
    """Temporary in-memory vector store for VSS document chunks using FAISS"""

    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        """
        Initialize the in-memory vector store

        Args:
            model_name: Sentence transformer model to use for embeddings
        """
        self.model_name = model_name
        self.embedding_model = SentenceTransformer(model_name)
        self.chunks: List[VSSChunk] = []
        self.embeddings: Optional[np.ndarray] = None
        self.index: Optional[faiss.IndexFlatIP] = None

    async def add_vss_documents(self, vss_paths: List[str]) -> None:
        """
        Add VSS documents to the vector store with exact page extraction

        Args:
            vss_paths: List of paths to VSS files (PDF/DOCX)
        """
        import asyncio
        loop = asyncio.get_event_loop()
        
        for file_path in vss_paths:
            try:
                ext = os.path.splitext(file_path)[1].lower()
                source_file = os.path.basename(file_path)

                if ext == ".pdf":
                    # Run file processing in thread executor to avoid blocking
                    await loop.run_in_executor(None, self._process_pdf, file_path, source_file)
                elif ext == ".docx":
                    # Run file processing in thread executor to avoid blocking
                    await loop.run_in_executor(None, self._process_docx, file_path, source_file)
                else:
                    logger.warning(f"Unsupported file type: {ext} for {file_path}")

            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                continue

        logger.info(f"Added {len(self.chunks)} chunks from {len(vss_paths)} VSS documents")

    def _process_pdf(self, file_path: str, source_file: str) -> None:
        """Process PDF file with exact page extraction"""
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            logger.info(f"Processing VSS PDF: {source_file}, Total pages: {total_pages}")

            for page_number, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if not page_text:
                    continue

                # Chunk the page text
                chunks = chunk_text(page_text)

                for chunk_index, chunk in enumerate(chunks):
                    vss_chunk = VSSChunk(
                        content=chunk.page_content,
                        page_number=page_number,
                        source_file=source_file,
                        chunk_id=f"{source_file}_page{page_number}_chunk{chunk_index}",
                        chunk_index=chunk_index,
                    )
                    self.chunks.append(vss_chunk)

    def _process_docx(self, file_path: str, source_file: str) -> None:
        """Process DOCX file - simulate pages by grouping paragraphs"""
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Group paragraphs into page-like chunks (~500 words per "page")
        current_page_text = []
        current_word_count = 0
        page_number = 1
        words_per_page = 500

        for paragraph in paragraphs:
            word_count = len(paragraph.split())

            if current_word_count + word_count > words_per_page and current_page_text:
                # Process current page
                page_text = "\n".join(current_page_text)
                chunks = chunk_text(page_text)

                for chunk_index, chunk in enumerate(chunks):
                    vss_chunk = VSSChunk(
                        content=chunk.page_content,
                        page_number=page_number,
                        source_file=source_file,
                        chunk_id=f"{source_file}_page{page_number}_chunk{chunk_index}",
                        chunk_index=chunk_index,
                    )
                    self.chunks.append(vss_chunk)

                # Start new page
                current_page_text = [paragraph]
                current_word_count = word_count
                page_number += 1
            else:
                current_page_text.append(paragraph)
                current_word_count += word_count

        # Process last page
        if current_page_text:
            page_text = "\n".join(current_page_text)
            chunks = chunk_text(page_text)

            for chunk_index, chunk in enumerate(chunks):
                vss_chunk = VSSChunk(
                    content=chunk.page_content,
                    page_number=page_number,
                    source_file=source_file,
                    chunk_id=f"{source_file}_page{page_number}_chunk{chunk_index}",
                    chunk_index=chunk_index,
                )
                self.chunks.append(vss_chunk)

        logger.info(f"Processed VSS DOCX: {source_file}, created {page_number} page chunks")

    async def build_index(self) -> None:
        """Build the FAISS index from all chunks asynchronously"""
        if not self.chunks:
            logger.warning("No chunks to index")
            return

        # Generate embeddings for all chunks using thread executor to avoid blocking
        texts = [chunk.content for chunk in self.chunks]
        logger.info(f"Generating embeddings for {len(texts)} VSS chunks...")
        
        import asyncio
        loop = asyncio.get_event_loop()
        
        # Run the CPU-intensive embedding generation in a thread executor
        self.embeddings = await loop.run_in_executor(
            None, self._generate_embeddings_sync, texts
        )

        # Create FAISS index
        dimension = self.embeddings.shape[1]
        self.index = faiss.IndexFlatIP(dimension)  # Inner product for cosine similarity

        # Normalize embeddings for cosine similarity
        faiss.normalize_L2(self.embeddings)
        self.index.add(self.embeddings.astype('float32'))

        logger.info(f"Built FAISS index with {len(self.chunks)} chunks, dimension {dimension}")
    
    def _generate_embeddings_sync(self, texts):
        """Synchronous helper method for embedding generation"""
        return self.embedding_model.encode(texts, show_progress_bar=True)

    def search_relevant_chunks(self, query: str, top_k: int = 5) -> List[VSSChunk]:
        """
        Search for relevant chunks using FAISS

        Args:
            query: Search query (indicator text)
            top_k: Number of top results to return

        Returns:
            List of relevant VSSChunk objects with similarity scores
        """
        if not self.index or self.embeddings is None:
            logger.warning("Index not built, returning empty results")
            return []

        # Encode query
        query_embedding = self.embedding_model.encode([query])
        faiss.normalize_L2(query_embedding)

        # Determine number of results to request; FAISS requires k > 0
        k = min(top_k, len(self.chunks))
        if k <= 0:
            logger.warning("FAISS search requested with k<=0; returning no results")
            return []

        # Search
        scores, indices = self.index.search(
            query_embedding.astype('float32'),
            k,
        )

        # Return chunks with scores
        results = []
        for score, idx in zip(scores[0], indices[0]):
            if idx < len(self.chunks):
                chunk = self.chunks[idx]
                # Add score to chunk for debugging
                chunk.similarity_score = float(score)
                results.append(chunk)

        logger.info(f"Found {len(results)} relevant VSS chunks for query: {query[:100]}...")
        return results

    def get_chunks_for_indicator(self, indicator_text: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """
        Get relevant chunks for an indicator in a format suitable for the prompt

        Args:
            indicator_text: The indicator text to search for
            top_k: Number of top chunks to return

        Returns:
            List of chunk dictionaries with metadata
        """
        chunks = self.search_relevant_chunks(indicator_text, top_k)

        formatted_chunks = []
        for chunk in chunks:
            formatted_chunk = {
                "content": chunk.content,
                "page_number": chunk.page_number,
                "source_file": chunk.source_file,
                "similarity_score": getattr(chunk, 'similarity_score', 0.0),
                "chunk_index": chunk.chunk_index,
            }
            formatted_chunks.append(formatted_chunk)

        return formatted_chunks

    def format_chunks_for_prompt(self, chunks: List[Dict[str, Any]]) -> str:
        """
        Format chunks into a string suitable for inclusion in the prompt

        Args:
            chunks: List of chunk dictionaries from get_chunks_for_indicator

        Returns:
            Formatted string with all chunks
        """
        if not chunks:
            return "No relevant VSS content found."

        formatted_parts = []
        for chunk in chunks:
            part = f"[Source: {chunk['source_file']}, Page: {chunk['page_number']}]\n{chunk['content'].strip()}"
            formatted_parts.append(part)

        return "\n\n".join(formatted_parts)

    def clear(self) -> None:
        """Clear all data from the vector store"""
        self.chunks.clear()
        self.embeddings = None
        self.index = None
        logger.info("Cleared in-memory VSS vector store")

    def get_stats(self) -> Dict[str, Any]:
        """Get statistics about the vector store"""
        return {
            "total_chunks": len(self.chunks),
            "index_built": self.index is not None,
            "embedding_dimension": self.embeddings.shape[1] if self.embeddings is not None else None,
            "source_files": list(set(chunk.source_file for chunk in self.chunks)),
            "total_pages": sum(1 for chunk in self.chunks if chunk.chunk_index == 0),  # Count first chunk of each page
        }


